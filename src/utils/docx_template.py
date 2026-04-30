"""DOCX Template Utilities — bank-agnostic helpers and analyst memo renderer.

Public functions:
  render_analyst_memo() — creates a new analyst memo DOCX with Output 2 + 3

Bank-specific form rendering (Output 1) lives in src/banks/{bank}/renderer.py.
Shared cell helpers (_set_cell, _append_cell, _fmt, etc.) are imported by
bank renderers from this module.

Cell coordinate reference: docs/requirements/docx_template_design.md
"""

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt

from ..utils.logger import get_logger

logger = get_logger("docx_template")


# ─────────────────────────────────────────────────────────────────────────────
# Public API
# ─────────────────────────────────────────────────────────────────────────────

def render_analyst_memo(
    output_path: str,
    company_name: str,
    section2: str | None = None,
    section3: str | None = None,
    date: str | None = None,
) -> str:
    """Create a credit analyst memo DOCX (tờ trình thẩm định tín dụng nội bộ).

    Creates a fresh Document (not from template) with:
      - Cover block: company name, date, classification
      - Output 2: sector analysis (from section2 markdown)
      - Output 3: financial analysis (from section3 markdown)

    Args:
        output_path:  Destination DOCX path.
        company_name: Company name for the header.
        section2:     Sector analysis markdown (Output 2 from subgraph2).
        section3:     Financial analysis markdown (Output 3 from subgraph3).
        date:         Report date. Defaults to today DD/MM/YYYY.

    Returns:
        output_path on success.
    """
    if date is None:
        date = datetime.now().strftime("%d/%m/%Y")

    doc = Document()

    # ── Cover block ─────────────────────────────────────────────────────────
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run("TỜ TRÌNH THẨM ĐỊNH TÍN DỤNG NỘI BỘ")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()  # blank spacer

    doc.add_paragraph(f"Khách hàng: {company_name}")
    doc.add_paragraph(f"Ngày lập: {date}")
    doc.add_paragraph("Phân loại: Nội bộ – Bảo mật")

    doc.add_paragraph()  # blank spacer

    # ── Output 2: Sector analysis ────────────────────────────────────────────
    if section2 and section2.strip():
        _append_markdown(doc, section2)
        logger.debug("Sector section appended to analyst memo")

    # ── Output 3: Financial analysis ─────────────────────────────────────────
    if section3 and section3.strip():
        _add_page_break(doc)
        _append_markdown(doc, section3)
        logger.debug("Financial section appended to analyst memo")

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    logger.info(f"Analyst memo DOCX saved → {output_path}")
    return output_path


# ─────────────────────────────────────────────────────────────────────────────
# Low-level cell helpers
# ─────────────────────────────────────────────────────────────────────────────

def _set_cell(cell, text: str | None) -> None:
    """Clear all runs in a cell and set new text, preserving paragraph style."""
    if text is None:
        return
    text = str(text).strip()
    if not text:
        return
    # Clear existing run content across all paragraphs
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    # Write into first paragraph
    if cell.paragraphs:
        para = cell.paragraphs[0]
        if para.runs:
            para.runs[0].text = text
        else:
            para.add_run(text)
    else:
        cell.add_paragraph(text)


def _append_cell(cell, text: str) -> None:
    """Append text to the last paragraph of a cell without clearing existing content."""
    if not text:
        return
    if cell.paragraphs:
        cell.paragraphs[-1].add_run(str(text))
    else:
        cell.add_paragraph(str(text))


# ─────────────────────────────────────────────────────────────────────────────
# Formatting helpers
# ─────────────────────────────────────────────────────────────────────────────

def _fmt(value: float | None, unit: str = "triệu đồng") -> str:
    """Format a financial value: None → '', ≥1000 → 'X,XXX.X tỷ đồng'."""
    if value is None or value == 0:
        return ""
    if abs(value) >= 1000:
        return f"{value / 1000:,.1f} tỷ đồng"
    return f"{value:,.1f} {unit}"


def _add_page_break(doc: Document) -> None:
    para = doc.add_paragraph()
    para.add_run().add_break(WD_BREAK.PAGE)


# ─────────────────────────────────────────────────────────────────────────────
# Markdown → DOCX append
# ─────────────────────────────────────────────────────────────────────────────

def _add_heading_safe(doc: Document, text: str, level: int) -> None:
    """Add a heading paragraph.

    python-docx's add_heading() uses BabelFish.ui2internal() which lowercases
    'Heading 1' → 'heading 1'. The template stores the style as 'Heading 1'
    (capital H), so lookup by name fails. Workaround: look up by style_id
    ('Heading1'), suppress the deprecation warning, then set text.
    """
    import warnings
    style_id = f"Heading{level}"
    para = doc.add_paragraph()
    try:
        with warnings.catch_warnings():
            warnings.simplefilter("ignore", UserWarning)
            para.style = doc.styles[style_id]
        para.add_run(text)
    except KeyError:
        # Fallback: bold paragraph if heading style not defined in template
        run = para.add_run(text)
        run.bold = True


def _append_markdown(doc: Document, md_text: str) -> None:
    """Parse markdown and append as DOCX elements (headings, tables, lists, text)."""
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith("# ") and not line.startswith("## "):
            _add_heading_safe(doc, line[2:].strip(), level=1)
            i += 1

        elif line.startswith("## ") and not line.startswith("### "):
            _add_heading_safe(doc, line[3:].strip(), level=2)
            i += 1

        elif line.startswith("### "):
            _add_heading_safe(doc, line[4:].strip(), level=3)
            i += 1

        elif line.startswith("|"):
            tbl_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                tbl_lines.append(lines[i])
                i += 1
            _md_table_to_docx(doc, tbl_lines)

        elif line.startswith("- ") or line.startswith("* "):
            # Template has no "List Bullet" style — use "List Paragraph" with bullet prefix
            para = doc.add_paragraph(style="List Paragraph")
            para.add_run("• " + line[2:].strip())
            i += 1

        elif re.match(r"^\d+\. ", line):
            # Template has no "List Number" style — keep number prefix in "List Paragraph"
            para = doc.add_paragraph(style="List Paragraph")
            para.add_run(line.strip())
            i += 1

        elif re.match(r"^-{3,}$", line.strip()):
            i += 1  # skip horizontal rules

        elif line.strip() == "":
            i += 1

        else:
            _add_paragraph_with_inline(doc, line.strip())
            i += 1


def _add_paragraph_with_inline(doc: Document, text: str) -> None:
    """Add a normal paragraph, parsing **bold** and *italic* inline markers."""
    para = doc.add_paragraph()
    pattern = r"(\*\*[^*]+\*\*|\*[^*]+\*)"
    for part in re.split(pattern, text):
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            para.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            para.add_run(part[1:-1]).italic = True
        else:
            para.add_run(part)


def _md_table_to_docx(doc: Document, table_lines: list[str]) -> None:
    """Convert markdown table lines into a DOCX Table Grid table."""
    rows = []
    for line in table_lines:
        if re.match(r"^\|[\s\-:|]+\|$", line.strip()):
            continue  # skip separator rows
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)

    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=num_cols)
    # Template only has "Table Normal" — look up by style_id "TableGrid" if available,
    # otherwise fall back to "Table Normal" (no visible borders, content still readable)
    import warnings
    try:
        with warnings.catch_warnings():
            warnings.simplefilter("ignore", UserWarning)
            tbl.style = doc.styles["TableGrid"]
    except KeyError:
        try:
            tbl.style = doc.styles["Table Normal"]
        except KeyError:
            pass

    for ri, row_data in enumerate(rows):
        for ci in range(num_cols):
            text = row_data[ci] if ci < len(row_data) else ""
            cell = tbl.cell(ri, ci)
            cell.text = ""
            run = cell.paragraphs[0].add_run(text)
            if ri == 0:
                run.bold = True
