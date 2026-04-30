"""VPBank form renderer — fills the giay-de-nghi-vay-von.docx template."""
from pathlib import Path

from docx import Document

from ...config import get_form_template_docx
from ...models.company import CompanyInfo
from ...models.financial import FinancialData
from ...utils.docx_template import _set_cell, _fmt, _append_cell
from ...utils.logger import get_logger

logger = get_logger("banks.vpbank.renderer")

_BANK = "vpbank"


def render(
    output_path: str,
    company_info: CompanyInfo | None = None,
    financial_data: FinancialData | None = None,
) -> str:
    """Fill VPBank DOCX template with real data and save to output_path.

    Output DOCX preserves template structure exactly (35 tables, 19 sections,
    13 images, green borders). No new sections are appended — Phụ lục A/B
    analysis text is only written to the markdown output, not to the DOCX.

    Args:
        output_path:    Destination DOCX file path.
        company_info:   Extracted company info (CompanyInfo model).
        financial_data: Extracted financial data (FinancialData model).

    Returns:
        output_path on success.
    """
    template_path = str(get_form_template_docx(_BANK))
    logger.info(f"Loading template: {template_path}")
    doc = Document(template_path)

    if company_info:
        _fill_company_info(doc, company_info)
        _fill_shareholders(doc, company_info)
        _fill_phu_luc_1(doc, company_info)
        _fill_board_and_management(doc, company_info)
        _fill_business_ops(doc, company_info)

    if financial_data and financial_data.statements:
        _fill_vot_thuc_gop(doc, financial_data)
        _fill_financial_history(doc, financial_data)
        _fill_income_statement(doc, financial_data)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    logger.info(f"DOCX saved → {output_path}")
    return output_path


# ─────────────────────────────────────────────────────────────────────────────
# Section fillers (VPBank-specific table layout)
# ─────────────────────────────────────────────────────────────────────────────

def _fill_company_info(doc: Document, info: CompanyInfo) -> None:
    """Fill Table[1] (1.1 legal entity) and Table[2] row 14 (representative).

    Table[1] layout: 13 rows × 2 cols — col 0 = label, col 1 = value cell.
    """
    t1 = doc.tables[1]

    _set_cell(t1.rows[2].cells[1], info.company_name)        # Tên Khách hàng
    # Tax code + ngày cấp + cơ quan cấp → right cell (cells[1])
    if info.tax_code:
        reg_value = info.tax_code
        if info.established_date:
            reg_value += f" Ngày cấp: {info.established_date}"
        if info.registration_authority:
            reg_value += f" Cơ quan cấp: {info.registration_authority}"
        _set_cell(t1.rows[3].cells[1], reg_value)
    _set_cell(t1.rows[4].cells[1], info.address)             # Địa chỉ trụ sở
    _set_cell(t1.rows[5].cells[1], info.address)             # Địa chỉ giao dịch (same)
    _set_cell(t1.rows[6].cells[1], info.phone)               # Điện thoại
    _set_cell(t1.rows[7].cells[1], info.main_business)       # Ngành nghề KD chính
    _set_cell(t1.rows[9].cells[1], info.charter_capital)     # Vốn điều lệ

    # Table[2]: row 14 col 1 = Người đại diện pháp luật
    if info.legal_representative:
        t2 = doc.tables[2]
        _set_cell(t2.rows[14].cells[1], info.legal_representative)

    logger.debug("Company info filled (Table[1] + Table[2] r14)")


def _fill_shareholders(doc: Document, info: CompanyInfo) -> None:
    """Fill Table[3] rows 10–12 — Section 1.3 cổ đông/thành viên góp vốn.

    Template has 3 pre-existing empty data rows (r10, r11, r12).
    Columns (9-col table, some merged): [0]=STT [1+2]=Họ tên [3+4]=Mối QH
    [5]=Tỷ lệ% [6+7]=Kinh nghiệm [8]=Dư nợ VPB.
    """
    if not info.shareholders:
        return

    t3 = doc.tables[3]
    for i, sh in enumerate(info.shareholders[:3]):
        row = t3.rows[10 + i]
        _set_cell(row.cells[0], str(i + 1))           # STT
        _set_cell(row.cells[1], sh.name)               # Họ và tên (c1+c2 merged)
        if sh.percentage is not None:
            _set_cell(row.cells[5], f"{sh.percentage:.1f}%")  # Tỷ lệ góp vốn

    logger.debug(f"Shareholders filled: {min(len(info.shareholders), 3)} rows")


def _fill_phu_luc_1(doc: Document, info: CompanyInfo) -> None:
    """Fill PHỤ LỤC 1 with the primary (largest) shareholder — individual section.

    PHỤ LỤC 1 is a FORM (label | value rows), not a data table.
    Template structure:
      Table[14] r0–r13 : Section Doanh nghiệp (skipped — MST shareholders are individuals)
      Table[14] r14    : Section header "Thành viên góp vốn chính (nếu là cá nhân)"
      Table[14] r15 c1 : Mối quan hệ với khách hàng
      Table[14] r16 c1 : Họ và tên
      Table[15] r5  c1 : Tỷ lệ góp vốn

    Fills only shareholders[0] (the primary/largest shareholder).
    """
    if not info.shareholders:
        return

    sh = info.shareholders[0]

    try:
        t14 = doc.tables[14]
        t15 = doc.tables[15]
    except IndexError:
        logger.warning("Table[14]/[15] not found — skipping PHỤ LỤC 1 fill")
        return

    # Table[14] r15 col1: Mối quan hệ với khách hàng → fixed "Cổ đông chính"
    if len(t14.rows) > 15:
        _set_cell(t14.rows[15].cells[1], "Cổ đông chính")

    # Table[14] r16 col1: Họ và tên
    if len(t14.rows) > 16:
        _set_cell(t14.rows[16].cells[1], sh.name)

    # Table[15] r5 col1: Tỷ lệ góp vốn
    if len(t15.rows) > 5 and sh.percentage is not None:
        _set_cell(t15.rows[5].cells[1], f"{sh.percentage:.1f}%")

    logger.debug(f"PHỤ LỤC 1 (Table[14]/[15]) filled for primary shareholder: {sh.name}")


def _fill_vot_thuc_gop(doc: Document, financial_data: FinancialData) -> None:
    """Fill Table[1] row 10 col 1 — Vốn thực góp đến ngày...

    human_mapping note #1: Vốn thực góp = Vốn chủ sở hữu (equity, mã 400 CĐKT).
    Uses the latest year available in financial_data.
    """
    years = sorted(financial_data.statements.keys())
    if not years:
        return
    s = financial_data.statements[years[-1]]
    if s.equity is None:
        return

    t1 = doc.tables[1]
    if len(t1.rows) > 10:
        _set_cell(t1.rows[10].cells[1], _fmt(s.equity))
        logger.debug(f"Vốn thực góp (Table[1] r10) filled: {_fmt(s.equity)} (equity {years[-1]})")


def _fill_board_and_management(doc: Document, info: CompanyInfo) -> None:
    """No-op: HĐQT/BGĐ/BKS are not part of the VPBank loan application form.

    Form section 1.2 only captures the legal representative (filled in
    _fill_company_info via Table[2] r14). Board/management data belongs
    in the analyst memo, not this customer-facing form.
    """
    logger.debug("Board/management injection skipped — not applicable to loan form")


def _fill_business_ops(doc: Document, info: CompanyInfo) -> None:
    """Fill Table[8] row 1 col 0 — Lĩnh vực kinh doanh chính."""
    if not info.main_business:
        return
    _set_cell(doc.tables[8].rows[1].cells[0], info.main_business)
    logger.debug("Business ops (Table[8]) filled")


def _fill_financial_history(doc: Document, financial_data: FinancialData) -> None:
    """Fill Table[29] — PHỤ LỤC 6: historical revenue / costs / net profit.

    Template: [0]=TT [1]=Chỉ tiêu [2]=Năm N-1 [3]=Năm kế hoạch [4]=Ghi chú
    Renderer renames col headers to actual years and fills with real data.
    Uses second-to-last year for col 2 and last year for col 3.
    """
    t29 = doc.tables[29]
    years = sorted(financial_data.statements.keys())
    if not years:
        return

    year_n1 = years[-2] if len(years) >= 2 else None
    year_n  = years[-1]

    # Rename column headers to actual years
    if year_n1:
        _set_cell(t29.rows[0].cells[2], str(year_n1))
    _set_cell(t29.rows[0].cells[3], str(year_n))

    for col_idx, year in [(2, year_n1), (3, year_n)]:
        if year is None:
            continue
        s = financial_data.statements.get(year)
        if not s:
            continue

        # Row 1: Doanh thu
        _set_cell(t29.rows[1].cells[col_idx], _fmt(s.net_revenue))

        # Row 2: Tổng chi phí = COGS + bán hàng + quản lý
        total_costs = sum(
            x for x in [s.cost_of_goods_sold, s.selling_expenses, s.admin_expenses]
            if x is not None
        )
        _set_cell(t29.rows[2].cells[col_idx], _fmt(total_costs or None))

        # Row 3: Lợi nhuận sau thuế
        _set_cell(t29.rows[3].cells[col_idx], _fmt(s.net_profit))

        # Row 4: Tổng nhu cầu vốn lưu động = Tài sản ngắn hạn - Nợ ngắn hạn (working capital)
        if s.current_assets is not None and s.current_liabilities is not None:
            working_capital = s.current_assets - s.current_liabilities
            _set_cell(t29.rows[4].cells[col_idx], _fmt(working_capital))

        # Row 5: Nguồn vốn tự có = Vốn chủ sở hữu (human_mapping note #1)
        _set_cell(t29.rows[5].cells[col_idx], _fmt(s.equity))

        # Row 6: Nhu cầu vốn vay TCTD khác = Tổng nợ phải trả (proxy, human_mapping note #2)
        _set_cell(t29.rows[6].cells[col_idx], _fmt(s.total_liabilities))

    logger.debug(f"Financial history (Table[29]) filled — years: {year_n1}, {year_n} (rows 1–6)")


def _fill_income_statement(doc: Document, financial_data: FinancialData) -> None:
    """Fill Table[32] — PHỤ LỤC 6: income statement detail for the latest year.

    Table[32] layout: [0]=STT [1]=Chi tiết [2]=12-month value [3]=Ghi chú
    Rows: 0=header 1=revenue 2=COGS 3=gross_profit 4=other_income
          5=finance_costs 6=tax 7=other_costs 8=net_profit
    """
    years = sorted(financial_data.statements.keys())
    if not years:
        return
    year = years[-1]
    s = financial_data.statements[year]

    t32 = doc.tables[32]
    _set_cell(t32.rows[0].cells[2], f"Năm {year}")          # Update header

    _set_cell(t32.rows[1].cells[2], _fmt(s.net_revenue))
    _set_cell(t32.rows[2].cells[2], _fmt(s.cost_of_goods_sold))
    _set_cell(t32.rows[3].cells[2], _fmt(s.gross_profit))
    # rows 4-6: other income / finance costs / tax — no data, leave blank
    other_costs = (s.admin_expenses or 0) + (s.selling_expenses or 0)
    _set_cell(t32.rows[7].cells[2], _fmt(other_costs or None))  # Chi phí khác
    _set_cell(t32.rows[8].cells[2], _fmt(s.net_profit))

    logger.debug(f"Income statement (Table[32]) filled for year {year}")
